import io

import pytest
from docx import Document

from app.integrations.text_extract import DOCX_MIME, extract_text


def test_extract_docx_round_trips_text():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Python engineer, 5 years")
    doc.save(buf)
    text = extract_text(DOCX_MIME, buf.getvalue())
    assert "Jane Doe" in text
    assert "Python engineer, 5 years" in text


def test_extract_text_rejects_unknown_type():
    with pytest.raises(ValueError):
        extract_text("text/plain", b"hello")
